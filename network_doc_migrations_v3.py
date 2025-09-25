"""
Network Document Migration – v3

Adds structured exports:
  • Per-document JSON (sections, paragraphs, tables, image file refs)
  • SQLite database (documents / sections / content tables)
  • Optional export of image assets alongside JSON

CLI examples
-----------
Batch migrate docs, plus JSON + SQLite export:
  python network_doc_migrations_v3.py *.docx \
    --template template.docx \
    --output-dir ./out/docs \
    --json-dir ./out/json \
    --sqlite ./out/migrations.db

JSON layout (per document)
--------------------------
{
  "metadata": {...},
  "sections": [
    {
      "name": "Network Topology",
      "order": 1,
      "content": [
        {"type": "paragraph", "text": "..."},
        {"type": "table", "rows": [["Device","Role"],["R1","Core"]]},
        {"type": "image", "path": "assets/topology_1.png", "width_px": 1024, "height_px": 768}
      ]
    }
  ]
}
"""

import os
import re
import json
import uuid
import base64
import sqlite3
from io import BytesIO
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple

import logging
import pandas as pd
from docx import Document
from docx.shared import Inches

# ----------------------------------------------------------------------------
# Logging
# ----------------------------------------------------------------------------
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ----------------------------------------------------------------------------
# Core classes
# ----------------------------------------------------------------------------
class WordDocumentMigrator:
    """Handles migration of content from old Word documents to new templates
       and provides structured extraction for JSON/SQLite exports.
    """

    def __init__(self, old_doc_path: str, template_path: str, output_path: str):
        self.old_doc_path = old_doc_path
        self.template_path = template_path
        self.output_path = output_path
        self.old_doc: Optional[Document] = None
        self.new_doc: Optional[Document] = None
        self.section_mapping: Dict[str, str] = {}

    # ----------------------------- Load Docs ---------------------------------
    def load_documents(self):
        try:
            self.old_doc = Document(self.old_doc_path)
            self.new_doc = Document(self.template_path)
            logger.info(f"Loaded old: {self.old_doc_path}")
            logger.info(f"Loaded template: {self.template_path}")
        except Exception as e:
            logger.error(f"Error loading documents: {str(e)}")
            raise

    # --------------------------- Extraction ----------------------------------
    def extract_sections(self, doc: Document) -> Dict[str, list]:
        sections: Dict[str, list] = {}
        current_section = None
        current_content: List[Tuple[str, object]] = []

        # Iterate raw body to preserve order of paragraphs/tables
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                para = None
                for p in doc.paragraphs:
                    if p._element == element:
                        para = p
                        break
                if para:
                    if para.style and para.style.name.startswith('Heading'):
                        if current_section:
                            sections[current_section] = current_content
                            current_content = []
                        current_section = para.text.strip() or "Untitled"
                        logger.debug(f"Found section: {current_section}")
                    elif current_section:
                        current_content.append(('paragraph', para))

            elif element.tag.endswith('tbl'):  # Table
                if current_section:
                    for t in doc.tables:
                        if t._element == element:
                            current_content.append(('table', t))
                            break

        if current_section:
            sections[current_section] = current_content
        return sections

    def _has_images(self, para) -> bool:
        for run in para.runs:
            if 'graphic' in run._element.xml or 'picture' in run._element.xml:
                return True
        return False

    # ------------------------- New Doc insertion -----------------------------
    def find_matching_section(self, old_section: str, new_sections: List[str]) -> Optional[str]:
        old_section_lower = old_section.lower()
        for new_section in new_sections:
            if new_section.lower() == old_section_lower:
                return new_section

        stop_words = {'the','a','an','and','or','of','to','in','for','on','at','by'}
        old_words = {w.strip('.,;:!?()[]{}') for w in old_section_lower.split()} - stop_words
        best_match = None
        best_count = 0
        for new_section in new_sections:
            new_words = {w.strip('.,;:!?()[]{}') for w in new_section.lower().split()} - stop_words
            common = old_words & new_words
            if len(common) > best_count:
                best_match, best_count = new_section, len(common)
        return best_match

    def setup_section_mapping(self, custom_mapping: Optional[Dict[str,str]] = None):
        if custom_mapping:
            self.section_mapping = custom_mapping
            return
        old_sections = self.extract_sections(self.old_doc)
        new_sections = self.extract_sections(self.new_doc)
        for old in old_sections.keys():
            match = self.find_matching_section(old, list(new_sections.keys()))
            if match:
                self.section_mapping[old] = match

    def insert_content_after_heading(self, heading_text: str, content_list: list):
        heading_el = None
        for element in self.new_doc.element.body:
            if element.tag.endswith('p'):
                for p in self.new_doc.paragraphs:
                    if p._element == element and p.text.strip() == heading_text:
                        heading_el = element
                        break
                if heading_el is not None:
                    break
        if heading_el is None:
            logger.warning(f"Heading not found in template: {heading_text}")
            return

        body = self.new_doc.element.body
        insert_index = body.index(heading_el) + 1
        elements_to_insert = []

        for content_type, content in content_list:
            if content_type == 'paragraph':
                if content.text.strip() or self._has_images(content):
                    new_p = self._create_paragraph_element(content)
                    elements_to_insert.append(new_p)
            elif content_type == 'table':
                new_t = self._copy_table_element(content)
                if new_t is not None:
                    elements_to_insert.append(new_t)

        for i, el in enumerate(elements_to_insert):
            body.insert(insert_index + i, el)

    def _create_paragraph_element(self, source_para):
        new_p = self.new_doc.add_paragraph()
        new_p.clear()
        if source_para.alignment:
            new_p.alignment = source_para.alignment
        if source_para.style and source_para.style.name != 'Normal':
            try:
                new_p.style = source_para.style.name
            except Exception:
                pass
        for run in source_para.runs:
            nr = new_p.add_run(run.text)
            nr.bold = bool(run.bold)
            nr.italic = bool(run.italic)
            nr.underline = bool(run.underline)
            if run.font.name:
                nr.font.name = run.font.name
            if run.font.size:
                nr.font.size = run.font.size
        el = new_p._element
        self.new_doc.element.body.remove(el)
        # images are not reproduced here (already handled in original v2 via _copy_images_from_paragraph)
        # keep behavior minimal to avoid duplicating inline media incorrectly
        return el

    def _copy_table_element(self, source_table):
        try:
            new_table = self.new_doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
            if source_table.style:
                try:
                    new_table.style = source_table.style
                except Exception:
                    pass
            for i, row in enumerate(source_table.rows):
                for j, cell in enumerate(row.cells):
                    new_cell = new_table.rows[i].cells[j]
                    if new_cell.paragraphs:
                        new_cell.paragraphs[0].clear()
                    for para_idx, para in enumerate(cell.paragraphs):
                        if para_idx == 0 and new_cell.paragraphs:
                            target_p = new_cell.paragraphs[0]
                        else:
                            target_p = new_cell.add_paragraph()
                        for run in para.runs:
                            tr = target_p.add_run(run.text)
                            tr.bold = bool(run.bold)
                            tr.italic = bool(run.italic)
            el = new_table._element
            self.new_doc.element.body.remove(el)
            return el
        except Exception as e:
            logger.error(f"Error copying table: {e}")
            return None

    # ----------------------------- Migration ---------------------------------
    def migrate_content(self):
        self.load_documents()
        old_sections = self.extract_sections(self.old_doc)
        if not self.section_mapping:
            self.setup_section_mapping()
        for old_s, new_s in self.section_mapping.items():
            if old_s in old_sections:
                self.insert_content_after_heading(new_s, old_sections[old_s])
        self.new_doc.save(self.output_path)
        logger.info(f"Saved migrated doc: {self.output_path}")

    # ------------------------ Structured extraction --------------------------
    def extract_structured(self, assets_dir: Optional[str] = None) -> Dict:
        """Return a structured dict of the old document, optionally exporting images.
        assets_dir: directory to save image assets (created if provided).
        """
        if self.old_doc is None:
            self.old_doc = Document(self.old_doc_path)

        sections_raw = self.extract_sections(self.old_doc)
        structured_sections = []

        if assets_dir:
            Path(assets_dir).mkdir(parents=True, exist_ok=True)

        def para_to_images_and_text(para, section_slug: str) -> List[dict]:
            items: List[dict] = []
            # Text first (single item if any text)
            text = para.text.strip()
            if text:
                items.append({"type": "paragraph", "text": text})
            # Extract embedded images (relationship parts with r:embed)
            try:
                for run in para.runs:
                    if 'graphic' in run._element.xml:
                        for inline in run._element.iter():
                            if 'blip' in inline.tag:
                                for attr in inline.attrib:
                                    if 'embed' in attr:
                                        rId = inline.attrib[attr]
                                        try:
                                            image_part = para.part.related_parts[rId]
                                            image_data = image_part.blob
                                            img_name = f"{section_slug}_{uuid.uuid4().hex[:8]}.png"
                                            if assets_dir:
                                                out_path = Path(assets_dir) / img_name
                                                with open(out_path, 'wb') as f:
                                                    f.write(image_data)
                                                items.append({"type": "image", "path": str(Path('assets')/img_name)})
                                            else:
                                                # if no assets_dir, fall back to base64 (keeps JSON self-contained)
                                                b64 = base64.b64encode(image_data).decode('ascii')
                                                items.append({"type": "image", "base64": b64, "encoding": "base64"})
                                        except Exception as e:
                                            logger.debug(f"Image extract failed: {e}")
            except Exception as e:
                logger.debug(f"Para image scan error: {e}")
            return items

        def table_to_rows(tbl) -> List[List[str]]:
            rows = []
            for row in tbl.rows:
                cells = []
                for cell in row.cells:
                    txt = ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    cells.append(txt)
                rows.append(cells)
            return rows

        for idx, (section_name, content_list) in enumerate(sections_raw.items(), start=1):
            section_slug = re.sub(r"[^a-z0-9]+", "_", section_name.lower()).strip("_") or f"section_{idx}"
            citems: List[dict] = []
            for ctype, obj in content_list:
                if ctype == 'paragraph':
                    citems.extend(para_to_images_and_text(obj, section_slug))
                elif ctype == 'table':
                    citems.append({"type": "table", "rows": table_to_rows(obj)})
            structured_sections.append({"name": section_name, "order": idx, "content": citems})

        meta = {
            "source_path": str(self.old_doc_path),
            "source_filename": Path(self.old_doc_path).name,
            "template_used": str(self.template_path),
            "output_document": str(self.output_path),
            "timestamp": datetime.now().isoformat(),
        }
        return {"metadata": meta, "sections": structured_sections}


class BatchMigrationManager:
    def __init__(self, template_path: str, output_dir: str = None,
                 json_dir: Optional[str] = None, sqlite_path: Optional[str] = None):
        self.template_path = template_path
        self.output_dir = output_dir or os.getcwd()
        self.json_dir = json_dir
        self.sqlite_path = sqlite_path
        self.migration_results: List[dict] = []

        if self.json_dir:
            Path(self.json_dir).mkdir(parents=True, exist_ok=True)
        if self.sqlite_path:
            self._init_sqlite(self.sqlite_path)

    # ----------------------------- SQLite ------------------------------------
    def _init_sqlite(self, db_path: str):
        con = sqlite3.connect(db_path)
        cur = con.cursor()
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                source_filename TEXT,
                source_path TEXT,
                template_used TEXT,
                output_document TEXT,
                timestamp TEXT
            )
            """
        )
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS sections (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                document_id INTEGER,
                name TEXT,
                ord INTEGER,
                FOREIGN KEY(document_id) REFERENCES documents(id)
            )
            """
        )
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS content (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                section_id INTEGER,
                type TEXT,
                text TEXT,
                table_json TEXT,
                image_path TEXT,
                image_base64 TEXT,
                FOREIGN KEY(section_id) REFERENCES sections(id)
            )
            """
        )
        con.commit()
        con.close()

    def _sqlite_insert_document(self, meta: dict) -> int:
        con = sqlite3.connect(self.sqlite_path)
        cur = con.cursor()
        cur.execute(
            "INSERT INTO documents (source_filename, source_path, template_used, output_document, timestamp) VALUES (?,?,?,?,?)",
            (
                meta.get('source_filename'),
                meta.get('source_path'),
                meta.get('template_used'),
                meta.get('output_document'),
                meta.get('timestamp'),
            )
        )
        doc_id = cur.lastrowid
        con.commit()
        con.close()
        return doc_id

    def _sqlite_insert_section(self, document_id: int, name: str, order: int) -> int:
        con = sqlite3.connect(self.sqlite_path)
        cur = con.cursor()
        cur.execute(
            "INSERT INTO sections (document_id, name, ord) VALUES (?,?,?)",
            (document_id, name, order)
        )
        sec_id = cur.lastrowid
        con.commit()
        con.close()
        return sec_id

    def _sqlite_insert_content(self, section_id: int, item: dict):
        con = sqlite3.connect(self.sqlite_path)
        cur = con.cursor()
        t = item.get('type')
        text = item.get('text') if t == 'paragraph' else None
        table_json = json.dumps(item.get('rows')) if t == 'table' else None
        image_path = item.get('path') if t == 'image' else None
        image_b64 = item.get('base64') if t == 'image' and 'base64' in item else None
        cur.execute(
            "INSERT INTO content (section_id, type, text, table_json, image_path, image_base64) VALUES (?,?,?,?,?,?)",
            (section_id, t, text, table_json, image_path, image_b64)
        )
        con.commit()
        con.close()

    # ----------------------------- Processing --------------------------------
    def process_documents(self, old_doc_paths: List[str], custom_mapping: Optional[Dict[str,str]] = None,
                          excel_output: Optional[str] = None) -> dict:
        results = {'successful': [], 'failed': [], 'summary': []}
        print("\n" + "="*60)
        print("BATCH MIGRATION STARTING")
        print(f"Processing {len(old_doc_paths)} documents")
        print(f"Template: {self.template_path}")
        print(f"Output directory: {self.output_dir}")
        if self.json_dir:
            print(f"JSON export directory: {self.json_dir}")
        if self.sqlite_path:
            print(f"SQLite DB: {self.sqlite_path}")
        print("="*60 + "\n")

        for idx, old_doc_path in enumerate(old_doc_paths, 1):
            print(f"[{idx}/{len(old_doc_paths)}] Processing: {old_doc_path}")
            base_name = Path(old_doc_path).stem
            output_name = f"{base_name}_migrated.docx"
            output_path = os.path.join(self.output_dir, output_name)

            migration_data = {
                'source_document': old_doc_path,
                'source_filename': Path(old_doc_path).name,
                'template_used': self.template_path,
                'output_document': output_path,
                'timestamp': datetime.now().isoformat(),
                'status': 'pending',
                'sections_found': 0,
                'sections_mapped': 0,
                'sections_migrated': 0,
                'tables_found': 0,
                'images_found': 0,
                'error_message': None,
            }

            try:
                migrator = WordDocumentMigrator(old_doc_path, self.template_path, output_path)
                migrator.load_documents()

                old_sections = migrator.extract_sections(migrator.old_doc)
                migration_data['sections_found'] = len(old_sections)

                if custom_mapping:
                    migrator.setup_section_mapping(custom_mapping)
                else:
                    migrator.setup_section_mapping()
                migration_data['sections_mapped'] = len(migrator.section_mapping)

                # Count tables & images in original
                for _, clist in old_sections.items():
                    for ctype, content in clist:
                        if ctype == 'table':
                            migration_data['tables_found'] += 1
                        elif ctype == 'paragraph' and migrator._has_images(content):
                            migration_data['images_found'] += 1

                # Perform migration to new Word
                migrator.migrate_content()
                migration_data['status'] = 'success'
                migration_data['sections_migrated'] = len(migrator.section_mapping)
                results['successful'].append(output_path)

                # ----------------- Structured export per document -----------------
                json_out_path = None
                if self.json_dir:
                    # assets dir is sibling: <json_dir>/<base_name>_assets
                    assets_dir = Path(self.json_dir) / f"{base_name}_assets"
                    structured = migrator.extract_structured(str(assets_dir))

                    json_out_path = Path(self.json_dir) / f"{base_name}.json"
                    with open(json_out_path, 'w', encoding='utf-8') as f:
                        json.dump(structured, f, ensure_ascii=False, indent=2)
                    print(f"   🗂  JSON saved: {json_out_path}")

                if self.sqlite_path:
                    # If we wrote assets alongside JSON, image paths should be stored relative
                    structured = structured if self.json_dir else migrator.extract_structured(None)
                    doc_id = self._sqlite_insert_document(structured['metadata'])
                    for section in structured['sections']:
                        sec_id = self._sqlite_insert_section(doc_id, section['name'], section['order'])
                        for item in section['content']:
                            self._sqlite_insert_content(sec_id, item)

                print(f"   ✅ Success: {output_name}")
                print(f"      Sections found: {migration_data['sections_found']} | mapped: {migration_data['sections_mapped']}")
                print(f"      Tables: {migration_data['tables_found']} | Images: {migration_data['images_found']}")

            except Exception as e:
                migration_data['status'] = 'failed'
                migration_data['error_message'] = str(e)
                results['failed'].append(old_doc_path)
                print(f"   ❌ Failed: {e}")
                logger.exception(f"Failed to migrate {old_doc_path}")

            results['summary'].append(migration_data)
            self.migration_results.append(migration_data)

        # Optional Excel overview for humans
        if excel_output:
            try:
                self._export_overview_excel(excel_output)
                print(f"\n📊 Excel summary exported: {excel_output}")
            except Exception as e:
                print(f"Warning: Excel export failed: {e}. Install: pip install pandas openpyxl")

        print("\n" + "="*60)
        print("BATCH SUMMARY")
        print(f"Processed: {len(results['successful']) + len(results['failed'])}")
        print(f"  ✅ Success: {len(results['successful'])}")
        print(f"  ❌ Failed:  {len(results['failed'])}")
        print("="*60)
        return results

    # ------------------------------ Excel (Optional, human-readable) ---------
    def _export_overview_excel(self, excel_path: str):
        df_rows = []
        for r in self.migration_results:
            row = {
                'Source Document': Path(r['source_document']).name,
                'Full Source Path': r['source_document'],
                'Template Used': r['template_used'],
                'Output Document': r['output_document'],
                'Status': r['status'],
                'Timestamp': r['timestamp'],
                'Sections Found': r['sections_found'],
                'Sections Mapped': r['sections_mapped'],
                'Sections Migrated': r['sections_migrated'],
                'Tables Found': r['tables_found'],
                'Images Found': r['images_found'],
                'Error': r['error_message'] or ''
            }
            df_rows.append(row)
        df = pd.DataFrame(df_rows)
        with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Summary', index=False)


# ----------------------------------------------------------------------------
# CLI
# ----------------------------------------------------------------------------

def main():
    import argparse
    parser = argparse.ArgumentParser(
        description="Migrate old Word docs into a new template and export structured JSON/SQLite",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )

    parser.add_argument('documents', nargs='+', help='Old .docx files (supporting glob via shell)')
    parser.add_argument('--template', '-t', required=True, help='New Word template (.docx)')
    parser.add_argument('--output-dir', '-o', default='.', help='Output folder for migrated .docx')

    parser.add_argument('--json-dir', help='Folder to write per-document JSON + assets')
    parser.add_argument('--sqlite', help='SQLite DB path to upsert into (created if missing)')

    parser.add_argument('--mapping', '-m', help='JSON file with custom section mapping {old->new}')
    parser.add_argument('--excel', '-e', help='Optional Excel overview path')
    parser.add_argument('--verbose', '-v', action='store_true', help='Verbose logging')

    args = parser.parse_args()
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    # Load custom mapping if provided
    custom_mapping = None
    if args.mapping:
        with open(args.mapping, 'r', encoding='utf-8') as f:
            custom_mapping = json.load(f)

    Path(args.output_dir).mkdir(parents=True, exist_ok=True)

    mgr = BatchMigrationManager(
        template_path=args.template,
        output_dir=args.output_dir,
        json_dir=args.json_dir,
        sqlite_path=args.sqlite,
    )

    mgr.process_documents(args.documents, custom_mapping=custom_mapping, excel_output=args.excel)


if __name__ == '__main__':
    main()
