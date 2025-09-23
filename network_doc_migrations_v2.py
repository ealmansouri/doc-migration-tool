class BatchMigrationManager:
    """Manages batch migration of multiple Word documents"""

    def __init__(self, template_path: str, output_dir: str = None):
        """
        Initialize batch migration manager

        Args:
            template_path: Path to the new Word template
            output_dir: Directory for output files (default: current directory)
        """
        self.template_path = template_path
        self.output_dir = output_dir or os.getcwd()
        self.migration_results = []

    def process_documents(self, old_doc_paths: list[str],
                          custom_mapping: dict[str, str] = None,
                          excel_output: str = None) -> dict:
        """
        Process multiple documents and optionally export to Excel

        Args:
            old_doc_paths: list of paths to old documents
            custom_mapping: Optional custom section mapping
            excel_output: Path for Excel summary (optional)

        Returns:
            Dictionary with migration results
        """
        results = {
            'successful': [],
            'failed': [],
            'summary': []
        }

        print(f"\n{'=' * 60}")
        print(f"BATCH MIGRATION STARTING")
        print(f"Processing {len(old_doc_paths)} documents")
        print(f"Template: {self.template_path}")
        print(f"Output directory: {self.output_dir}")
        print(f"{'=' * 60}\n")

        for idx, old_doc_path in enumerate(old_doc_paths, 1):
            print(f"\n[{idx}/{len(old_doc_paths)}] Processing: {old_doc_path}")
            print("-" * 40)

            # Generate output filename
            base_name = Path(old_doc_path).stem
            output_name = f"{base_name}_migrated.docx"
            output_path = os.path.join(self.output_dir, output_name)

            # Track migration data
            migration_data = {
                'source_document': old_doc_path,
                'source_filename': Path(old_doc_path).name,
                'template_used': self.template_path,
                'output_document': output_path,
                'timestamp': datetime.now().isoformat(),
                'status': 'pending',
                'sections_found': 0,
                'sections_mapped': 0,
                'sections_migrated': 0,
                'tables_migrated': 0,
                'images_migrated': 0,
                'error_message': None,
                'section_mapping': {},
                'original_sections': {}  # Store original document sections
            }

            try:
                # Create migrator for this document
                migrator = WordDocumentMigrator(old_doc_path, self.template_path, output_path)
                migrator.load_documents()

                # Extract ALL content from the old document, including unmapped sections
                old_sections = migrator.extract_sections(migrator.old_doc)
                migration_data['sections_found'] = len(old_sections)

                # Store ALL original sections with their content
                migration_data['original_sections'] = old_sections

                # Setup mapping
                if custom_mapping:
                    migrator.setup_section_mapping(custom_mapping)
                else:
                    migrator.setup_section_mapping()

                migration_data['sections_mapped'] = len(migrator.section_mapping)
                migration_data['section_mapping'] = migrator.section_mapping

                # Count content types in original document
                for section_name, content_list in old_sections.items():
                    for content_type, content in content_list:
                        if content_type == 'table':
                            migration_data['tables_migrated'] += 1
                        elif content_type == 'paragraph' and migrator._has_images(content):
                            migration_data['images_migrated'] += 1

                # Perform migration
                migrator.migrate_content()

                migration_data['status'] = 'success'
                migration_data['sections_migrated'] = len(migrator.section_mapping)

                results['successful'].append(output_path)
                print(f"✅ Success: {output_name}")
                print(f"   Original sections extracted: {migration_data['sections_found']}")
                print(f"   Sections mapped to template: {migration_data['sections_mapped']}")
                print(f"   Tables found: {migration_data['tables_migrated']}")
                print(f"   Images found: {migration_data['images_migrated']}")

            except Exception as e:
                migration_data['status'] = 'failed'
                migration_data['error_message'] = str(e)
                results['failed'].append(old_doc_path)
                print(f"❌ Failed: {str(e)}")
                logger.error(f"Failed to migrate {old_doc_path}: {str(e)}")

            # Add to results
            results['summary'].append(migration_data)
            self.migration_results.append(migration_data)

        # Print summary
        self._print_summary(results)

        # Export to Excel if requested
        if excel_output:
            self._export_to_excel(excel_output)
            print(f"\n📊 Excel summary exported to: {excel_output}")

        return results

    def _print_summary(self, results: dict):
        """Print migration summary"""
        print(f"\n{'=' * 60}")
        print("BATCH MIGRATION SUMMARY")
        print(f"{'=' * 60}")
        print(f"Total processed: {len(results['successful']) + len(results['failed'])}")
        print(f"✅ Successful: {len(results['successful'])}")
        print(f"❌ Failed: {len(results['failed'])}")

        if results['failed']:
            print("\nFailed documents:")
            for doc in results['failed']:
                print(f"  - {doc}")

    def _export_to_excel(self, excel_path: str):
        """
        Export original document content to Excel spreadsheet with each section as a column

        Args:
            excel_path: Path for Excel output file
        """
        if not self.migration_results:
            logger.warning("No migration results to export")
            return

        # Prepare data for DataFrame with original document content
        excel_data = []
        all_original_sections = set()

        # First pass: collect all unique section names from original documents
        for result in self.migration_results:
            if 'original_sections' in result:
                for section_name in result['original_sections'].keys():
                    all_original_sections.add(section_name)

        # Sort section names for consistent column order
        all_original_sections = sorted(list(all_original_sections))

        print(f"\nFound {len(all_original_sections)} unique sections across all documents:")
        for i, section in enumerate(all_original_sections, 1):
            print(f"  {i}. {section}")

        # Second pass: build rows with metadata and original section content
        for result in self.migration_results:
            # Start with metadata
            row = {
                'Source Document': result['source_filename'],
                'Full Source Path': result['source_document'],
                'Template Used': result['template_used'],
                'Output Document': result['output_document'],
                'Migration Status': result['status'],
                'Timestamp': result['timestamp'],
                'Total Sections in Original': result['sections_found'],
                'Sections Mapped': result['sections_mapped'],
                'Sections Migrated': result['sections_migrated'],
                'Total Tables': result['tables_migrated'],
                'Total Images': result['images_migrated'],
                'Error Message': result['error_message'] or ''
            }

            # Add original section content columns
            if 'original_sections' in result:
                for section_name in all_original_sections:
                    column_name = f'[OLD] {section_name}'
                    if section_name in result['original_sections']:
                        # Extract text content from the original section
                        content = result['original_sections'][section_name]
                        content_text = self._extract_text_from_content(content)

                        # Excel has a 32,767 character limit per cell
                        if len(content_text) > 32000:
                            content_text = content_text[:31997] + "..."
                        row[column_name] = content_text
                    else:
                        row[column_name] = ''
            else:
                # No content available, fill with empty strings
                for section_name in all_original_sections:
                    column_name = f'[OLD] {section_name}'
                    row[column_name] = ''

            # Add mapping information
            if result['section_mapping']:
                mapping_str = ' | '.join([f"{old} → {new}"
                                          for old, new in result['section_mapping'].items()])
                row['Section Mappings Applied'] = mapping_str
            else:
                row['Section Mappings Applied'] = 'No mappings'

            excel_data.append(row)

        # Create DataFrame
        df = pd.DataFrame(excel_data)

        # Reorder columns: metadata first, then original sections
        metadata_cols = ['Source Document', 'Full Source Path', 'Template Used',
                         'Output Document', 'Migration Status', 'Timestamp',
                         'Total Sections in Original', 'Sections Mapped', 'Sections Migrated',
                         'Total Tables', 'Total Images', 'Section Mappings Applied', 'Error Message']
        section_cols = [col for col in df.columns if col.startswith('[OLD]')]
        ordered_cols = metadata_cols + section_cols
        df = df[ordered_cols]

        # Export with formatting
        try:
            with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
                # Main sheet with all original document data
                df.to_excel(writer, sheet_name='Original Document Content', index=False)

                # Create summary sheet
                summary_df = df[metadata_cols].copy()
                summary_df.to_excel(writer, sheet_name='Migration Summary', index=False)

                # Create section index sheet
                section_index_data = []
                for idx, section in enumerate(all_original_sections, 1):
                    docs_with_section = 0
                    for result in self.migration_results:
                        if 'original_sections' in result and section in result['original_sections']:
                            docs_with_section += 1

                    section_index_data.append({
                        'Section #': idx,
                        'Section Name': section,
                        'Found in # Documents': docs_with_section,
                        'Found in % Documents': f"{(docs_with_section / len(self.migration_results) * 100):.1f}%"
                    })

                section_index_df = pd.DataFrame(section_index_data)
                section_index_df.to_excel(writer, sheet_name='Section Index', index=False)

                # Format main sheet
                workbook = writer.book
                worksheet = writer.sheets['Original Document Content']

                # Auto-adjust column widths with reasonable limits
                for column in df:
                    column_length = max(df[column].astype(str).map(len).max(), len(column))
                    col_idx = df.columns.get_loc(column)

                    # Convert column index to Excel column letter(s)
                    if col_idx < 26:
                        col_letter = chr(65 + col_idx)
                    else:
                        first_letter = chr(65 + (col_idx // 26) - 1)
                        second_letter = chr(65 + (col_idx % 26))
                        col_letter = first_letter + second_letter

                    # Set width with limits
                    if column.startswith('[OLD]'):
                        # Original content columns - wider but with max
                        worksheet.column_dimensions[col_letter].width = min(column_length + 2, 70)
                    else:
                        # Metadata columns - normal width
                        worksheet.column_dimensions[col_letter].width = min(column_length + 2, 40)

                # Add color formatting and styles
                from openpyxl.styles import PatternFill, Font, Alignment, Border, Side

                # Define colors
                green_fill = PatternFill(start_color='90EE90', end_color='90EE90', fill_type='solid')
                red_fill = PatternFill(start_color='FFB6C1', end_color='FFB6C1', fill_type='solid')
                header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
                section_header_fill = PatternFill(start_color='70AD47', end_color='70AD47', fill_type='solid')
                header_font = Font(color='FFFFFF', bold=True, size=11)
                border = Border(bottom=Side(style='thin'))

                # Format headers
                for col_idx, cell in enumerate(worksheet[1], 1):
                    if cell.value and cell.value.startswith('[OLD]'):
                        cell.fill = section_header_fill
                    else:
                        cell.fill = header_fill
                    cell.font = header_font
                    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    cell.border = border

                # Set row height for header
                worksheet.row_dimensions[1].height = 30

                # Format status column
                status_col = df.columns.get_loc('Migration Status') + 1
                for row_num, status in enumerate(df['Migration Status'], start=2):
                    cell = worksheet.cell(row=row_num, column=status_col)
                    if status == 'success':
                        cell.fill = green_fill
                    elif status == 'failed':
                        cell.fill = red_fill

                # Enable text wrapping for content cells
                for col_idx, col_name in enumerate(df.columns, start=1):
                    if col_name.startswith('[OLD]'):
                        for row_idx in range(2, len(df) + 2):
                            cell = worksheet.cell(row=row_idx, column=col_idx)
                            cell.alignment = Alignment(wrap_text=True, vertical='top', horizontal='left')

                # Freeze panes (freeze first row and first three columns)
                worksheet.freeze_panes = 'D2'

                # Format summary sheet
                summary_worksheet = writer.sheets['Migration Summary']
                for cell in summary_worksheet[1]:
                    cell.fill = header_fill
                    cell.font = header_font
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                    cell.border = border

                # Auto-adjust summary sheet columns
                for column in summary_df:
                    column_length = max(summary_df[column].astype(str).map(len).max(), len(column))
                    col_idx = summary_df.columns.get_loc(column)
                    if col_idx < 26:
                        col_letter = chr(65 + col_idx)
                    else:
                        col_letter = chr(65 + (col_idx // 26) - 1) + chr(65 + (col_idx % 26))
                    summary_worksheet.column_dimensions[col_letter].width = min(column_length + 2, 50)

                # Format section index sheet
                index_worksheet = writer.sheets['Section Index']
                for cell in index_worksheet[1]:
                    cell.fill = header_fill
                    cell.font = header_font
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                    cell.border = border

                # Auto-adjust section index columns
                for column in section_index_df:
                    column_length = max(section_index_df[column].astype(str).map(len).max(), len(column))
                    col_idx = section_index_df.columns.get_loc(column)
                    col_letter = chr(65 + col_idx)
                    index_worksheet.column_dimensions[col_letter].width = column_length + 2

            logger.info(f"Excel file with original document content exported to: {excel_path}")
            print(f"\n📊 Excel file created with:")
            print(f"   - {len(all_original_sections)} unique section columns from original documents")
            print(f"   - {len(self.migration_results)} document rows")
            print(f"   - 3 sheets: Original Content, Migration Summary, Section Index")

        except Exception as e:
            logger.error(f"Failed to export Excel: {str(e)}")
            print(f"Warning: Could not export to Excel: {str(e)}")
            print("Make sure you have pandas and openpyxl installed: pip install pandas openpyxl")

    def _extract_text_from_content(self, content_list):
        """
        Extract text from content list (paragraphs and tables)

        Args:
            content_list: list of (type, content) tuples

        Returns:
            Combined text string
        """
        text_parts = []

        for content_type, content in content_list:
            if content_type == 'paragraph':
                # Extract paragraph text
                para_text = content.text.strip()
                if para_text:
                    text_parts.append(para_text)
            elif content_type == 'table':
                # Extract table text
                table_text = self._extract_table_text(content)
                if table_text:
                    text_parts.append(f"[TABLE]\n{table_text}")

        return '\n\n'.join(text_parts)

    def _extract_table_text(self, table):
        """
        Extract text content from a table

        Args:
            table: Word table object

        Returns:
            Formatted table text
        """
        rows_text = []

        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = ' '.join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                row_cells.append(cell_text)
            rows_text.append(' | '.join(row_cells))

        return '\n'.join(rows_text)
        """
Word Document Content Migration Script
Migrates content from old HLD/LLD Word documents to a new Word template
Preserves text, tables, and images/diagrams
Supports batch processing and Excel export
"""


import os
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from copy import deepcopy
import logging
from typing import Dict, List, Optional, Tuple
from io import BytesIO
from datetime import datetime
import pandas as pd
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class WordDocumentMigrator:
    """Handles migration of content from old Word documents to new templates"""

    def __init__(self, old_doc_path: str, template_path: str, output_path: str):
        """
        Initialize the migrator with document paths

        Args:
            old_doc_path: Path to the old Word document
            template_path: Path to the new Word template
            output_path: Path for the output document
        """
        self.old_doc_path = old_doc_path
        self.template_path = template_path
        self.output_path = output_path
        self.old_doc = None
        self.new_doc = None
        self.section_mapping = {}

    def load_documents(self):
        """Load the old document and new template"""
        try:
            self.old_doc = Document(self.old_doc_path)
            self.new_doc = Document(self.template_path)
            logger.info(f"Successfully loaded old document: {self.old_doc_path}")
            logger.info(f"Successfully loaded template: {self.template_path}")
        except Exception as e:
            logger.error(f"Error loading documents: {str(e)}")
            raise

    def extract_sections(self, doc: Document) -> dict[str, list]:
        """
        Extract sections and their content from a document

        Args:
            doc: The Word document to extract from

        Returns:
            Dictionary with section headings as keys and content lists as values
        """
        sections = {}
        current_section = None
        current_content = []

        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                para = None
                for p in doc.paragraphs:
                    if p._element == element:
                        para = p
                        break

                if para:
                    # Check if this is a heading
                    if para.style.name.startswith('Heading'):
                        # Save previous section if exists
                        if current_section:
                            sections[current_section] = current_content
                            current_content = []
                        current_section = para.text.strip()
                        logger.info(f"Found section: {current_section}")
                    elif current_section:
                        # Add paragraph to current section
                        current_content.append(('paragraph', para))

            elif element.tag.endswith('tbl'):  # Table
                if current_section:
                    for t in doc.tables:
                        if t._element == element:
                            current_content.append(('table', t))
                            logger.info(f"Found table in section: {current_section}")
                            break

        # Save last section
        if current_section:
            sections[current_section] = current_content

        return sections

    def find_matching_section(self, old_section: str, new_sections: list[str]) -> Optional[str]:
        """
        Find the best matching section in the new document

        Args:
            old_section: Section heading from old document
            new_sections: list of section headings in new document

        Returns:
            Best matching section heading or None
        """
        old_section_lower = old_section.lower()

        # First try exact match
        for new_section in new_sections:
            if new_section.lower() == old_section_lower:
                return new_section

        # Clean and split into words, removing common words that don't help matching
        stop_words = {'the', 'a', 'an', 'and', 'or', 'of', 'to', 'in', 'for', 'on', 'at', 'by'}

        # Get meaningful words from old section
        old_words = set(old_section_lower.split()) - stop_words
        old_words = {word.strip('.,;:!?()[]{}') for word in old_words if word.strip('.,;:!?()[]{})')}

        best_match = None
        best_match_count = 0

        # Try to find section with most words in common
        for new_section in new_sections:
            new_section_lower = new_section.lower()
            # Get meaningful words from new section
            new_words = set(new_section_lower.split()) - stop_words
            new_words = {word.strip('.,;:!?()[]{}') for word in new_words if word.strip('.,;:!?()[]{})')}

            # Find common words
            common_words = old_words & new_words

            # If at least 1 word matches, consider it a potential match
            if len(common_words) >= 1:
                if len(common_words) > best_match_count:
                    best_match = new_section
                    best_match_count = len(common_words)

        return best_match

    def insert_content_after_heading(self, heading_text: str, content_list: list[Tuple]):
        """
        Insert content after a specific heading in the new document

        Args:
            heading_text: The heading to insert content after
            content_list: list of (type, content) tuples
        """
        # Find the heading in the document element structure
        heading_element = None
        for element in self.new_doc.element.body:
            if element.tag.endswith('p'):
                for p in self.new_doc.paragraphs:
                    if p._element == element and p.text.strip() == heading_text:
                        heading_element = element
                        break
                if heading_element is not None:
                    break

        if heading_element is None:
            logger.warning(f"Could not find heading: {heading_text}")
            return

        # Get the parent (body) and the index of the heading
        body = self.new_doc.element.body
        heading_index = body.index(heading_element)
        insert_index = heading_index + 1

        # Track elements we're adding so we can insert them in order
        elements_to_insert = []

        for content_type, content in content_list:
            if content_type == 'paragraph':
                if content.text.strip() or self._has_images(content):  # Include if has text or images
                    # Create new paragraph with same content
                    new_para = self._create_paragraph_element(content)
                    elements_to_insert.append(new_para)

            elif content_type == 'table':
                # Create new table element
                new_table = self._copy_table_element(content)
                if new_table is not None:
                    elements_to_insert.append(new_table)

        # Insert all elements at the correct position
        for i, element in enumerate(elements_to_insert):
            body.insert(insert_index + i, element)

        logger.info(f"Inserted {len(elements_to_insert)} elements after section: {heading_text}")

    def _create_paragraph_element(self, source_para):
        """
        Create a new paragraph element with content from source paragraph

        Args:
            source_para: Source paragraph to copy

        Returns:
            New paragraph element
        """
        from docx.oxml import parse_xml
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph

        # Create new paragraph through the document
        new_p = self.new_doc.add_paragraph()

        # Clear default text
        new_p.clear()

        # Copy alignment if exists
        if source_para.alignment:
            new_p.alignment = source_para.alignment

        # Copy paragraph style if applicable
        if source_para.style and source_para.style.name != 'Normal':
            try:
                new_p.style = source_para.style.name
            except:
                pass

        # Copy runs with formatting
        for run in source_para.runs:
            new_run = new_p.add_run(run.text)

            # Copy formatting
            if run.bold:
                new_run.bold = True
            if run.italic:
                new_run.italic = True
            if run.underline:
                new_run.underline = True
            if run.font.name:
                new_run.font.name = run.font.name
            if run.font.size:
                new_run.font.size = run.font.size

        # Extract the element and remove from document body (we'll insert it at right place)
        element = new_p._element
        self.new_doc.element.body.remove(element)

        # Check for and copy images
        self._copy_images_from_paragraph(source_para, new_p)

        return element

    def _copy_table_element(self, source_table):
        """
        Create a copy of a table element

        Args:
            source_table: Source table to copy

        Returns:
            New table element
        """
        try:
            # Create new table with same dimensions
            new_table = self.new_doc.add_table(rows=len(source_table.rows),
                                               cols=len(source_table.columns))

            # Copy table style if available
            if source_table.style:
                try:
                    new_table.style = source_table.style
                except:
                    logger.debug("Could not copy table style")

            # Copy cell content
            for i, row in enumerate(source_table.rows):
                for j, cell in enumerate(row.cells):
                    new_cell = new_table.rows[i].cells[j]

                    # Clear default paragraph
                    if new_cell.paragraphs:
                        p = new_cell.paragraphs[0]
                        p.clear()

                    # Copy all paragraphs from source cell
                    for para_idx, para in enumerate(cell.paragraphs):
                        if para_idx == 0 and new_cell.paragraphs:
                            # Use existing first paragraph
                            new_para = new_cell.paragraphs[0]
                        else:
                            new_para = new_cell.add_paragraph()

                        # Copy text and formatting
                        for run in para.runs:
                            new_run = new_para.add_run(run.text)
                            if run.bold:
                                new_run.bold = True
                            if run.italic:
                                new_run.italic = True

            # Extract element and remove from body
            element = new_table._element
            self.new_doc.element.body.remove(element)

            return element

        except Exception as e:
            logger.error(f"Error copying table: {str(e)}")
            return None

    def _has_images(self, para):
        """
        Check if a paragraph contains images

        Args:
            para: Paragraph to check

        Returns:
            Boolean indicating if paragraph has images
        """
        for run in para.runs:
            if 'graphic' in run._element.xml or 'picture' in run._element.xml:
                return True
        return False

    def _copy_images_from_paragraph(self, source_para, target_para):
        """
        Copy images from source paragraph to target paragraph

        Args:
            source_para: Source paragraph with potential images
            target_para: Target paragraph to add images to
        """
        try:
            # Check each run for inline shapes/images
            for run in source_para.runs:
                if 'graphic' in run._element.xml:
                    # Try to extract and copy the image
                    for inline in run._element.iter():
                        if 'blip' in inline.tag:
                            for attr in inline.attrib:
                                if 'embed' in attr:
                                    rId = inline.attrib[attr]
                                    try:
                                        image_part = source_para.part.related_parts[rId]
                                        image_data = image_part.blob

                                        # Add image to target paragraph
                                        new_run = target_para.add_run()
                                        new_run.add_picture(BytesIO(image_data), width=Inches(4))
                                        logger.debug("Successfully copied an image")
                                    except Exception as e:
                                        logger.warning(f"Could not copy image: {str(e)}")
        except Exception as e:
            logger.debug(f"Error checking for images: {str(e)}")

    def setup_section_mapping(self, custom_mapping: dict[str, str] = None):
        """
        Setup mapping between old and new document sections

        Args:
            custom_mapping: Optional custom mapping dictionary
        """
        if custom_mapping:
            self.section_mapping = custom_mapping
            print("\n=== CUSTOM SECTION MAPPING APPLIED ===")
            for old_section, new_section in custom_mapping.items():
                print(f"  '{old_section}' -> '{new_section}'")
        else:
            # Auto-detect mappings
            old_sections = self.extract_sections(self.old_doc)
            new_sections = self.extract_sections(self.new_doc)

            print("\n=== AUTO-DETECTING SECTION MAPPINGS ===")
            print("Looking for sections with at least 1 word in common...\n")

            # Define stop words to ignore in matching
            stop_words = {'the', 'a', 'an', 'and', 'or', 'of', 'to', 'in', 'for', 'on', 'at', 'by'}

            for old_section in old_sections.keys():
                matching = self.find_matching_section(old_section, list(new_sections.keys()))
                if matching:
                    self.section_mapping[old_section] = matching

                    # Show which words matched
                    old_words = set(old_section.lower().split()) - stop_words
                    old_words = {word.strip('.,;:!?()[]{}') for word in old_words if word.strip('.,;:!?()[]{})')}
                    new_words = set(matching.lower().split()) - stop_words
                    new_words = {word.strip('.,;:!?()[]{}') for word in new_words if word.strip('.,;:!?()[]{})')}
                    common_words = old_words & new_words

                    print(f"✓ MAPPED: '{old_section}'")
                    print(f"      TO: '{matching}'")
                    print(f"  COMMON: {', '.join(common_words)}")
                    print()

                    logger.info(f"Mapped: '{old_section}' -> '{matching}' (common words: {common_words})")
                else:
                    print(f"✗ NO MATCH: '{old_section}'")
                    print()
                    logger.warning(f"No matching section found for: '{old_section}'")

            # Summary
            print(f"\n=== MAPPING SUMMARY ===")
            print(f"Total sections in old document: {len(old_sections)}")
            print(f"Successfully mapped: {len(self.section_mapping)}")
            print(f"Unmapped sections: {len(old_sections) - len(self.section_mapping)}")

            if self.section_mapping:
                print("\n=== FINAL MAPPING TABLE ===")
                for old, new in self.section_mapping.items():
                    print(f"  '{old}' -> '{new}'")

    def migrate_content(self):
        """Main method to perform the content migration"""
        logger.info("Starting content migration...")

        # Load documents
        self.load_documents()

        # Extract sections from old document
        old_sections = self.extract_sections(self.old_doc)
        logger.info(f"Found {len(old_sections)} sections in old document")

        # Setup section mappings if not provided
        if not self.section_mapping:
            self.setup_section_mapping()

        # Migrate content based on mapping
        for old_section, new_section in self.section_mapping.items():
            if old_section in old_sections:
                content = old_sections[old_section]
                self.insert_content_after_heading(new_section, content)
            else:
                logger.warning(f"Section '{old_section}' not found in old document")

        # Save the new document
        self.new_doc.save(self.output_path)
        logger.info(f"Migration complete. Output saved to: {self.output_path}")

    def print_section_summary(self):
        """Print a summary of sections found in both documents"""
        if not self.old_doc or not self.new_doc:
            self.load_documents()

        old_sections = self.extract_sections(self.old_doc)
        new_sections = self.extract_sections(self.new_doc)

        print("\n=== OLD DOCUMENT SECTIONS ===")
        for i, section in enumerate(old_sections.keys(), 1):
            content_count = len(old_sections[section])
            print(f"{i}. {section} ({content_count} items)")

        print("\n=== NEW TEMPLATE SECTIONS ===")
        for i, section in enumerate(new_sections.keys(), 1):
            print(f"{i}. {section}")

        print("\n=== SECTION MAPPING ===")
        if self.section_mapping:
            for old, new in self.section_mapping.items():
                print(f"'{old}' -> '{new}'")
        else:
            print("No mapping defined yet")

    def _export_to_excel(self, excel_path: str):
        """
        Export migration results to Excel spreadsheet

        Args:
            excel_path: Path for Excel output file
        """
        if not self.migration_results:
            logger.warning("No migration results to export")
            return

        # Prepare data for DataFrame
        excel_data = []

        for result in self.migration_results:
            # Flatten section mapping for Excel
            section_mapping_str = '; '.join([f"{old}->{new}"
                                             for old, new in result['section_mapping'].items()])

            row = {
                'Source Document': result['source_filename'],
                'Full Source Path': result['source_document'],
                'Template Used': result['template_used'],
                'Output Document': result['output_document'],
                'Migration Status': result['status'],
                'Timestamp': result['timestamp'],
                'Sections Found': result['sections_found'],
                'Sections Mapped': result['sections_mapped'],
                'Sections Migrated': result['sections_migrated'],
                'Tables Migrated': result['tables_migrated'],
                'Images Migrated': result['images_migrated'],
                'Section Mappings': section_mapping_str,
                'Error Message': result['error_message'] or ''
            }
            excel_data.append(row)

        # Create DataFrame and export
        df = pd.DataFrame(excel_data)

        # Export with formatting
        with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Migration Summary', index=False)

            # Get workbook and worksheet
            workbook = writer.book
            worksheet = writer.sheets['Migration Summary']

            # Auto-adjust column widths
            for column in df:
                column_length = max(df[column].astype(str).map(len).max(), len(column))
                col_idx = df.columns.get_loc(column)
                worksheet.column_dimensions[chr(65 + col_idx)].width = min(column_length + 2, 50)

            # Add color formatting for status
            from openpyxl.styles import PatternFill
            green_fill = PatternFill(start_color='90EE90', end_color='90EE90', fill_type='solid')
            red_fill = PatternFill(start_color='FFB6C1', end_color='FFB6C1', fill_type='solid')

            status_col = df.columns.get_loc('Migration Status') + 1
            for row_num, status in enumerate(df['Migration Status'], start=2):
                cell = worksheet.cell(row=row_num, column=status_col)
                if status == 'success':
                    cell.fill = green_fill
                elif status == 'failed':
                    cell.fill = red_fill

        logger.info(f"Excel summary exported to: {excel_path}")


def main():
    """Main function to run the Word Document Migrator with command line arguments"""

    import argparse
    import sys
    import json

    # Set up argument parser
    parser = argparse.ArgumentParser(
        description="Migrate content from old Word HLD/LLD documents to new templates",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  Single document migration:
    python migrate_word.py old_doc.docx template.docx output.docx

  Multiple documents (batch mode):
    python migrate_word.py doc1.docx doc2.docx doc3.docx --template template.docx --output-dir ./output

  Batch mode with Excel export:
    python migrate_word.py *.docx --template template.docx --output-dir ./output --excel summary.xlsx

  With analysis only (no migration):
    python migrate_word.py old_doc.docx template.docx output.docx --analyze

  With custom mapping file:
    python migrate_word.py old_doc.docx template.docx output.docx --mapping mapping.json
        """
    )

    # Arguments for both single and batch modes
    parser.add_argument('documents',
                        nargs='+',
                        help='Path(s) to old Word document(s). For batch mode, provide multiple files')

    # Single mode arguments (positional, optional)
    parser.add_argument('template_single',
                        nargs='?',
                        help='Template path (for single document mode)')
    parser.add_argument('output_single',
                        nargs='?',
                        help='Output path (for single document mode)')

    # Batch mode arguments
    parser.add_argument('--template', '-t',
                        help='Path to the Word template (required for batch mode)')
    parser.add_argument('--output-dir', '-o',
                        help='Output directory for migrated documents (batch mode)')
    parser.add_argument('--excel', '-e',
                        help='Export migration summary to Excel file (batch mode)')

    # Optional arguments
    parser.add_argument('--analyze', '-a',
                        action='store_true',
                        help='Only analyze documents and show sections without migrating')
    parser.add_argument('--mapping', '-m',
                        help='Path to JSON file containing custom section mappings')
    parser.add_argument('--interactive', '-i',
                        action='store_true',
                        help='Review auto-detected mappings before migration')
    parser.add_argument('--verbose', '-v',
                        action='store_true',
                        help='Enable verbose logging')

    # Parse arguments
    args = parser.parse_args()

    # Configure logging level
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    # Determine if we're in batch mode or single mode
    batch_mode = len(args.documents) > 1 or args.template or args.output_dir or args.excel

    # Validate arguments based on mode
    if batch_mode:
        if not args.template:
            print("Error: --template is required for batch mode")
            sys.exit(1)
        template_path = args.template
        output_dir = args.output_dir or os.getcwd()
    else:
        # Single mode
        if not args.template_single or not args.output_single:
            print("Error: For single document mode, provide: old_doc template output")
            print("       For batch mode, use --template flag")
            sys.exit(1)
        template_path = args.template_single
        output_path = args.output_single

    # Load custom mapping if provided
    custom_mapping = {}
    if args.mapping:
        try:
            with open(args.mapping, 'r') as f:
                custom_mapping = json.load(f)
                print(f"Loaded custom mapping from: {args.mapping}")
        except FileNotFoundError:
            print(f"Error: Mapping file not found: {args.mapping}")
            sys.exit(1)
        except json.JSONDecodeError:
            print(f"Error: Invalid JSON in mapping file: {args.mapping}")
            sys.exit(1)

    try:
        if batch_mode:
            # Batch processing mode
            print(f"\n🔄 BATCH PROCESSING MODE")
            print(f"Documents to process: {len(args.documents)}")

            # Create output directory if it doesn't exist
            if args.output_dir:
                os.makedirs(args.output_dir, exist_ok=True)

            # Create batch manager
            batch_manager = BatchMigrationManager(template_path, output_dir)

            # Process documents
            results = batch_manager.process_documents(
                args.documents,
                custom_mapping,
                args.excel
            )

            # Exit with appropriate code
            if results['failed']:
                sys.exit(1)  # Exit with error if any migrations failed

        else:
            # Single document mode (original functionality)
            old_doc_path = args.documents[0]

            # Create migrator instance
            migrator = WordDocumentMigrator(old_doc_path, template_path, output_path)

            # Load documents
            migrator.load_documents()

            # Show section analysis
            print("\nAnalyzing documents...")
            migrator.print_section_summary()

            # If analyze only mode, exit here
            if args.analyze:
                print("\n[Analysis complete - no migration performed]")
                sys.exit(0)

            # Setup mapping
            if custom_mapping:
                migrator.setup_section_mapping(custom_mapping)
            else:
                # Auto-detect mappings
                migrator.setup_section_mapping()

            # Interactive mode - allow user to review mappings
            if args.interactive and not custom_mapping:
                print("\n" + "=" * 50)
                response = input("Do you want to proceed with these mappings? (yes/no/edit): ").lower().strip()

                if response == 'edit':
                    print("\nEnter custom mappings (or press Enter to skip):")
                    print("Format: old_section -> new_section")
                    print("Type 'done' when finished\n")

                    custom_mapping = {}
                    while True:
                        mapping_input = input("Mapping: ").strip()
                        if mapping_input.lower() == 'done':
                            break
                        if ' -> ' in mapping_input:
                            old, new = mapping_input.split(' -> ', 1)
                            custom_mapping[old.strip()] = new.strip()
                            print(f"  Added: {old.strip()} -> {new.strip()}")
                        elif mapping_input:
                            print("  Invalid format. Use: old_section -> new_section")

                    if custom_mapping:
                        migrator.setup_section_mapping(custom_mapping)

                elif response != 'yes':
                    print("Migration cancelled.")
                    sys.exit(0)

            # Perform migration
            print("\nStarting migration...")
            migrator.migrate_content()

            print(f"\n✅ Success! Migrated content saved to: {output_path}")

            # Option to save the mapping for future use
            if not custom_mapping and migrator.section_mapping:
                save_mapping = input("\nSave these mappings for future use? (yes/no): ").lower().strip()
                if save_mapping == 'yes':
                    mapping_file = output_path.replace('.docx', '_mapping.json')
                    with open(mapping_file, 'w') as f:
                        json.dump(migrator.section_mapping, f, indent=2)
                    print(f"Mappings saved to: {mapping_file}")

    except FileNotFoundError as e:
        print(f"\n❌ Error: Could not find file - {e}")
        sys.exit(1)
    except PermissionError as e:
        print(f"\n❌ Error: Permission denied - {e}")
        print("Make sure the files are not open in Word and you have write permissions.")
        sys.exit(1)
    except Exception as e:
        print(f"\n❌ Error during migration: {e}")
        if args.verbose:
            logger.error(f"Migration failed: {str(e)}", exc_info=True)
        sys.exit(1)


if __name__ == "__main__":
    main()